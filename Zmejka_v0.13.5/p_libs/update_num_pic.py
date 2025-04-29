"""
Автор оригинального скрипта: Равшан Ибатулин
Ремейк: Ниса Дипова
"""

import re
from docx import Document  # pip install python-docx
import tkinter as tk
from tkinter import filedialog

def number_figures_in_document(doc_path: str,
                               figure_first_count: int = 1,
                               table_first_count: int = 1):
    """Функция для сквозной нумерации рисунков в документе."""
    # Открываем документ
    doc = Document(doc_path)
    figure_count = figure_first_count  # Счетчик для рисунков
    num_table = 0  # Счетчик для таблиц
    num_para = 0
    amount_pictures = []

    # Проходим по всем элементам тела документа
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Если это параграф
            para = doc.paragraphs[num_para]
            # Ищем подписи к рисункам
            matches = re.findall(r'Рисунок \d+\.|Рис\. \d+\.', para.text)
            for match in matches:
                new_text = re.sub(r'Рисунок \d+\.|Рис\. \d+\.', f'Рисунок {figure_count}.', match)
                para.text = para.text.replace(match, new_text)
                figure_count += 1
            num_para += 1

        elif element.tag.endswith('tbl'):  # Если это таблица
            # Проходим по всем таблицам
            table = doc.tables[num_table]
            num_table += 1
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'Рисунок \d+\.|Рис\. \d+\.', cell.text)
                    for match in matches:
                        new_text = re.sub(r'Рисунок \d+\.|Рис\. \d+\.', f'Рисунок {figure_count}.', match)
                        cell.text = cell.text.replace(match, new_text)
                        figure_count += 1
    
    print(f'Общее количество рисунков: {figure_count - 1}')
    # Сохраняем изменения в документе
    doc.save('Откорректированный_' + doc_path.split('/')[-1])  # Сохраняем с новым именем

# Создаем корневое окно (не отображаем его)
root = tk.Tk()
root.withdraw()  # Скрываем главное окно

# Запрашиваем у пользователя путь к документу
report_name = filedialog.askopenfilename(title="Выберите отчет", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))

# Убедимся, что файл был выбран
if report_name:
    number_figures_in_document(doc_path=report_name)
else:
    print("Файл не выбран.")